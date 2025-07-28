from fpdf import FPDF
from docx import Document
from pathlib import Path

def generate_pdf(title: str, meta: str, output_path: Path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=14)
    
    pdf.cell(200, 10, txt="SEO Metadata Report", ln=True, align='C')
    pdf.ln(10)

    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(40, 10, txt="Title:")
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=title)
    
    pdf.ln(5)
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(60, 10, txt="Meta Description:")
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=meta)

    pdf.output(str(output_path))


def generate_docx(title: str, meta: str, output_path: Path):
    doc = Document()
    doc.add_heading("SEO Metadata Report", level=1)

    doc.add_heading("Title:", level=2)
    doc.add_paragraph(title)

    doc.add_heading("Meta Description:", level=2)
    doc.add_paragraph(meta)

    doc.save(str(output_path))
